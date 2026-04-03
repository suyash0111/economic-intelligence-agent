"""
Document Generator - Creates professionally formatted Word documents.
New 10-section architecture inspired by WEF, Goldman Sachs, and McKinsey.

Structure:
  1. Executive Brief (The One-Pager)
  2. Macro Pulse: Global Economic Dashboard
  3. Policy & Central Bank Watch
  4. Thematic Deep Dives
  5. Cross-Source Intelligence
  6. Regional Outlook
  7. Risk Radar
  8. Strategic Implications & Recommendations
  9. The Week Ahead
  10. Appendix: Full Article Index
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path
from datetime import datetime
import logging
import re

from collectors.base_collector import Article
from config.settings import Settings

logger = logging.getLogger(__name__)

# =========================================================================
# DESIGN CONSTANTS
# =========================================================================

# Color Palette
NAVY = RGBColor(0x0D, 0x1B, 0x2A)       # Primary headers
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x4B)   # Secondary headers
ACCENT_BLUE = RGBColor(0x00, 0x66, 0x99)  # Subheadings
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)    # Hyperlinks
GREY_TEXT = RGBColor(0x71, 0x80, 0x96)     # Captions, dates
LIGHT_GREY = RGBColor(0xA0, 0xAE, 0xC0)   # Subtle text
RED = RGBColor(0xB4, 0x00, 0x00)           # Critical items
ORANGE = RGBColor(0xC8, 0x64, 0x00)        # Important items
GREEN = RGBColor(0x00, 0x78, 0x00)         # Positive/notable
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Font
FONT_FAMILY = 'Calibri'
BODY_SIZE = Pt(10.5)
SMALL_SIZE = Pt(9)
CAPTION_SIZE = Pt(8.5)

# Hex color strings for XML shading
SHADE_TITLE_BG = "0D1B2A"
SHADE_SECTION_BG = "E8EDF2"
SHADE_ARTICLE_BG = "F4F6F8"
SHADE_CRITICAL = "FFF0F0"
SHADE_IMPORTANT = "FFF8F0"
SHADE_LIGHT_BLUE = "E8F4FD"
SHADE_LIGHT_GREEN = "E8F5E9"
SHADE_LIGHT_YELLOW = "FFFDE7"


class DocumentGenerator:
    """
    Generates professionally formatted Word documents for the
    Global Pulse Weekly Economic Intelligence Report.

    New 10-section architecture.
    """

    def __init__(self):
        self.doc = None

    # =====================================================================
    # PUBLIC API
    # =====================================================================

    def generate(
        self,
        articles_by_org: dict[str, list[Article]],
        executive_summary: str,
        date_range: str,
        output_path: Path = None,
        tldr_top5: str = "",
        cross_source_synthesis: str = "",
        theme_summary: dict = None,
        sentiment_analysis: str = "",
        actionable_implications: str = "",
        geographic_summary: dict = None,
        key_numbers: str = "",
        policy_watch: str = "",
        risk_radar: str = "",
        week_ahead: str = ""
    ) -> Path:
        """Generate a comprehensive Word document report."""
        self.doc = Document()
        self._setup_document()

        # Flatten all articles for appendix
        all_articles = []
        for org_articles in articles_by_org.values():
            all_articles.extend(org_articles)

        # === COVER PAGE ===
        self._add_cover_page(date_range)

        # === TABLE OF CONTENTS ===
        self._add_table_of_contents()

        # === SECTION 1: EXECUTIVE BRIEF ===
        self._add_section_header("EXECUTIVE BRIEF", "1")
        self._add_executive_brief(
            executive_summary=executive_summary,
            tldr_top5=tldr_top5,
            sentiment_analysis=sentiment_analysis,
            key_numbers=key_numbers
        )
        self.doc.add_page_break()

        # === SECTION 2: MACRO PULSE ===
        self._add_section_header("MACRO PULSE: GLOBAL ECONOMIC DASHBOARD", "2")
        self._add_macro_pulse(key_numbers)
        self.doc.add_page_break()

        # === SECTION 3: POLICY & CENTRAL BANK WATCH ===
        self._add_section_header("POLICY & CENTRAL BANK WATCH", "3")
        self._add_policy_watch(policy_watch, all_articles)
        self.doc.add_page_break()

        # === SECTION 4: THEMATIC DEEP DIVES ===
        self._add_section_header("THEMATIC DEEP DIVES", "4")
        self._add_thematic_deep_dives(all_articles, theme_summary)
        self.doc.add_page_break()

        # === SECTION 5: CROSS-SOURCE INTELLIGENCE ===
        if cross_source_synthesis:
            self._add_section_header("CROSS-SOURCE INTELLIGENCE", "5")
            self._add_intro_text("Analysis of consensus and divergent views across organizations.")
            self._add_formatted_text(cross_source_synthesis)
            self.doc.add_page_break()

        # === SECTION 6: REGIONAL OUTLOOK ===
        if geographic_summary:
            self._add_section_header("REGIONAL OUTLOOK", "6")
            self._add_regional_outlook(geographic_summary)
            self.doc.add_page_break()

        # === SECTION 7: RISK RADAR ===
        self._add_section_header("RISK RADAR", "7")
        self._add_risk_radar(risk_radar, all_articles)
        self.doc.add_page_break()

        # === SECTION 8: STRATEGIC IMPLICATIONS & RECOMMENDATIONS ===
        if actionable_implications:
            self._add_section_header("STRATEGIC IMPLICATIONS & RECOMMENDATIONS", "8")
            self._add_intro_text("Practical considerations for key stakeholders.")
            self._add_formatted_text(actionable_implications)
            self.doc.add_page_break()

        # === SECTION 9: THE WEEK AHEAD ===
        self._add_section_header("THE WEEK AHEAD", "9")
        self._add_week_ahead(week_ahead)
        self.doc.add_page_break()

        # === SECTION 10: APPENDIX ===
        self._add_section_header("APPENDIX: FULL ARTICLE INDEX", "10")
        self._add_article_appendix(all_articles, articles_by_org)

        # === METHODOLOGY & DISCLAIMER ===
        self._add_methodology_note()

        # Save
        if output_path is None:
            Settings.ensure_output_dir()
            output_path = Settings.OUTPUT_DIR / f"Global_Pulse_Weekly_Report_{datetime.now().strftime('%Y%m%d')}.docx"

        self.doc.save(output_path)
        logger.info(f"Document saved: {output_path}")
        return output_path

    # =====================================================================
    # DOCUMENT SETUP
    # =====================================================================

    def _setup_document(self):
        """Set up document defaults: font, margins, styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_FAMILY
        font.size = BODY_SIZE
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(2)
        pf.line_spacing = 1.15

        for level in range(1, 4):
            heading_style = self.doc.styles[f'Heading {level}']
            heading_style.font.name = FONT_FAMILY
            heading_style.font.color.rgb = NAVY

            if level == 1:
                heading_style.font.size = Pt(18)
                heading_style.paragraph_format.space_before = Pt(24)
                heading_style.paragraph_format.space_after = Pt(10)
            elif level == 2:
                heading_style.font.size = Pt(14)
                heading_style.paragraph_format.space_before = Pt(18)
                heading_style.paragraph_format.space_after = Pt(8)
            elif level == 3:
                heading_style.font.size = Pt(12)
                heading_style.font.color.rgb = ACCENT_BLUE
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)

        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        self._add_page_numbers()

    def _add_page_numbers(self):
        """Add page numbers to footer."""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run("The Global Pulse  |  ")
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_GREY
        run.font.name = FONT_FAMILY

        run2 = p.add_run("Page ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = LIGHT_GREY
        run2.font.name = FONT_FAMILY

        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>0</w:t></w:r></w:fldSimple>'
        ).format(nsdecls('w'))
        fld_elem = parse_xml(fld_xml)
        p._p.append(fld_elem)

    # =====================================================================
    # COVER PAGE
    # =====================================================================

    def _add_cover_page(self, date_range: str):
        """Add a professional cover page."""
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("THE GLOBAL PULSE")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = FONT_FAMILY
        title.paragraph_format.space_after = Pt(4)

        # Decorative line
        line = self.doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run("_" * 40)
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(12)
        line.paragraph_format.space_after = Pt(8)

        # Subtitle
        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run("Weekly Economic Intelligence Briefing")
        run.font.size = Pt(16)
        run.font.color.rgb = DARK_BLUE
        run.font.name = FONT_FAMILY
        sub.paragraph_format.space_after = Pt(20)

        # Date range
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(date_range)
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = GREY_TEXT
        run.font.name = FONT_FAMILY
        date_p.paragraph_format.space_after = Pt(40)

        # Description
        desc = self.doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for text in [
            "360-Degree Analysis of Global Economic Developments",
            "Covering 39 Leading Organizations Worldwide",
            "",
            "Macro & Micro Indicators  |  Policy Watch  |  Risk Radar  |  Strategic Recommendations",
        ]:
            run = desc.add_run(text + "\n")
            run.font.size = Pt(11)
            run.font.color.rgb = GREY_TEXT
            run.font.name = FONT_FAMILY

        # Powered by
        powered = self.doc.add_paragraph()
        powered.alignment = WD_ALIGN_PARAGRAPH.CENTER
        powered.paragraph_format.space_before = Pt(60)
        run = powered.add_run("Powered by NVIDIA NIM 7-Model Architecture")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = LIGHT_GREY

        self.doc.add_page_break()

    # =====================================================================
    # TABLE OF CONTENTS
    # =====================================================================

    def _add_table_of_contents(self):
        """Add a Table of Contents page."""
        h = self.doc.add_heading("TABLE OF CONTENTS", level=1)

        toc_items = [
            ("1.", "Executive Brief"),
            ("2.", "Macro Pulse: Global Economic Dashboard"),
            ("3.", "Policy & Central Bank Watch"),
            ("4.", "Thematic Deep Dives"),
            ("5.", "Cross-Source Intelligence"),
            ("6.", "Regional Outlook"),
            ("7.", "Risk Radar"),
            ("8.", "Strategic Implications & Recommendations"),
            ("9.", "The Week Ahead"),
            ("10.", "Appendix: Full Article Index"),
        ]

        for num, title in toc_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.left_indent = Inches(0.3)

            num_run = p.add_run(f"{num}  ")
            num_run.font.bold = True
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = ACCENT_BLUE

            title_run = p.add_run(title)
            title_run.font.size = Pt(11)

        self.doc.add_page_break()

    # =====================================================================
    # SECTION HEADERS & FORMATTING HELPERS
    # =====================================================================

    def _add_section_header(self, title: str, number: str = ""):
        """Add a visually distinct section header with background shading."""
        display_title = f"{number}. {title}" if number else title
        h = self.doc.add_heading(display_title, level=1)

        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SHADE_SECTION_BG}"/>')
        h.paragraph_format.element.get_or_add_pPr().append(shading)

        line = self.doc.add_paragraph()
        line.paragraph_format.space_after = Pt(8)
        line.paragraph_format.space_before = Pt(0)

    def _add_intro_text(self, text: str):
        """Add italic intro/description text below a section header."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = GREY_TEXT
        run.font.size = SMALL_SIZE
        p.paragraph_format.space_after = Pt(10)

    def _add_sub_heading(self, title: str):
        """Add a colored sub-heading within a section."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE

    # =====================================================================
    # SECTION 1: EXECUTIVE BRIEF
    # =====================================================================

    def _add_executive_brief(self, executive_summary: str, tldr_top5: str,
                             sentiment_analysis: str, key_numbers: str):
        """Build the one-pager executive brief."""
        self._add_intro_text("THE BOTTOM LINE — Your one-page weekly economic intelligence digest.")

        # --- Headline Verdict (Executive Summary) ---
        if executive_summary:
            self._add_sub_heading("HEADLINE VERDICT")
            self._add_formatted_text(executive_summary)

        # --- TL;DR Top 5 ---
        if tldr_top5:
            self._add_sub_heading("TOP 5 DEVELOPMENTS")
            self._add_tldr_content(tldr_top5)

        # --- Sentiment Gauge ---
        if sentiment_analysis:
            self._add_sub_heading("MARKET SENTIMENT")
            self._add_sentiment_content(sentiment_analysis)

        # --- Key Numbers Box ---
        if key_numbers:
            self._add_sub_heading("KEY NUMBERS AT A GLANCE")
            self._add_formatted_text(key_numbers)

    # =====================================================================
    # SECTION 2: MACRO PULSE
    # =====================================================================

    def _add_macro_pulse(self, key_numbers: str):
        """Add the economic dashboard section."""
        self._add_intro_text("Key economic indicators and data points for the reporting period.")

        if key_numbers:
            self._add_formatted_text(key_numbers)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Economic indicator data unavailable. Check FRED API connectivity.")
            run.font.italic = True
            run.font.color.rgb = GREY_TEXT

    # =====================================================================
    # SECTION 3: POLICY & CENTRAL BANK WATCH
    # =====================================================================

    def _add_policy_watch(self, policy_watch: str, articles: list[Article]):
        """Add policy and central bank section."""
        self._add_intro_text("Monetary policy decisions, central bank communications, and fiscal policy developments.")

        if policy_watch:
            self._add_formatted_text(policy_watch)
        else:
            # Auto-generate from central bank articles
            cb_articles = [a for a in articles if a.category == 'Central Bank']
            if cb_articles:
                self._add_sub_heading("CENTRAL BANK ACTIVITY")
                for article in sorted(cb_articles, key=lambda x: x.importance_score, reverse=True)[:8]:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_after = Pt(4)

                    source_run = p.add_run(f"[{article.source}] ")
                    source_run.font.bold = True
                    source_run.font.color.rgb = ACCENT_BLUE
                    source_run.font.size = BODY_SIZE

                    title_run = p.add_run(article.title)
                    title_run.font.size = BODY_SIZE

                    if article.ai_summary:
                        summary_p = self.doc.add_paragraph()
                        summary_p.paragraph_format.left_indent = Inches(0.4)
                        summary_p.paragraph_format.space_after = Pt(6)
                        run = summary_p.add_run(article.ai_summary)
                        run.font.size = SMALL_SIZE
                        run.font.color.rgb = GREY_TEXT
            else:
                p = self.doc.add_paragraph()
                run = p.add_run("No central bank publications detected this reporting period.")
                run.font.italic = True
                run.font.color.rgb = GREY_TEXT

    # =====================================================================
    # SECTION 4: THEMATIC DEEP DIVES
    # =====================================================================

    def _add_thematic_deep_dives(self, articles: list[Article], theme_summary: dict = None):
        """Add 2-3 deep dives into the week's most important reports."""
        self._add_intro_text("In-depth analysis of the week's most significant reports and developments.")

        # Find articles with deep analysis
        deep_articles = [a for a in articles if a.deep_analysis and a.deep_analysis.get('success')]

        if deep_articles:
            for i, article in enumerate(deep_articles[:3], 1):
                self._add_sub_heading(f"DEEP DIVE {i}: {article.title[:70]}")

                # Source and date
                meta_p = self.doc.add_paragraph()
                run = meta_p.add_run(f"Source: {article.source_full}")
                run.font.size = SMALL_SIZE
                run.font.color.rgb = GREY_TEXT
                if article.published_date:
                    run2 = meta_p.add_run(f"  |  {article.published_date.strftime('%B %d, %Y')}")
                    run2.font.size = SMALL_SIZE
                    run2.font.color.rgb = GREY_TEXT
                meta_p.paragraph_format.space_after = Pt(6)

                # Summary
                if article.ai_summary:
                    self._add_formatted_text(article.ai_summary)

                # Deep analysis content
                if article.ai_analysis:
                    self._add_formatted_text(article.ai_analysis)

                # Key statistics
                deep = article.deep_analysis
                if deep.get('key_statistics'):
                    stats_head = self.doc.add_paragraph()
                    run = stats_head.add_run("KEY STATISTICS")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = ACCENT_BLUE

                    for stat in deep['key_statistics'][:8]:
                        sp = self.doc.add_paragraph()
                        sp.paragraph_format.left_indent = Inches(0.3)
                        sp.paragraph_format.space_after = Pt(2)
                        run = sp.add_run(f"• {stat}")
                        run.font.size = SMALL_SIZE

                # Chart descriptions
                if deep.get('chart_descriptions'):
                    chart_head = self.doc.add_paragraph()
                    run = chart_head.add_run("CHARTS & VISUALS")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = ACCENT_BLUE

                    for j, desc in enumerate(deep['chart_descriptions'][:3], 1):
                        cp = self.doc.add_paragraph()
                        cp.paragraph_format.left_indent = Inches(0.3)
                        cp.paragraph_format.space_after = Pt(4)
                        label = cp.add_run(f"Chart {j}: ")
                        label.font.bold = True
                        label.font.size = SMALL_SIZE
                        text = cp.add_run(desc)
                        text.font.size = SMALL_SIZE

                # Table summaries
                if deep.get('table_summaries'):
                    for table_md in deep['table_summaries'][:3]:
                        self._render_markdown_table(table_md)

        elif theme_summary:
            # Fallback: use theme summary
            self._add_theme_content(theme_summary)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("No major reports with deep analysis available this week.")
            run.font.italic = True
            run.font.color.rgb = GREY_TEXT

    # =====================================================================
    # SECTION 6: REGIONAL OUTLOOK
    # =====================================================================

    def _add_regional_outlook(self, geographic: dict):
        """Add geographic breakdown as a compact table + summaries."""
        self._add_intro_text("Economic pulse by geographic region.")

        regions_with_data = {k: v for k, v in geographic.items() if v.get('count', 0) > 0}
        if not regions_with_data:
            return

        # Summary table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['Region', 'Articles', 'Sources', 'Key Headlines']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = WHITE
                    run.font.size = Pt(10)
                    run.font.name = FONT_FAMILY
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SHADE_TITLE_BG}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        for region, data in regions_with_data.items():
            row = table.add_row()
            row.cells[0].text = region
            row.cells[1].text = str(data.get('count', 0))
            row.cells[2].text = ", ".join(data.get('sources', [])[:4])
            headlines = data.get('headlines', [])[:2]
            row.cells[3].text = "\n".join([f"• {h[:60]}..." for h in headlines]) if headlines else "—"

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = FONT_FAMILY

        widths = [Inches(1.2), Inches(0.8), Inches(2.0), Inches(3.0)]
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = width

        self.doc.add_paragraph()  # Spacer

    # =====================================================================
    # SECTION 7: RISK RADAR
    # =====================================================================

    def _add_risk_radar(self, risk_radar: str, articles: list[Article]):
        """Add forward-looking risk assessment."""
        self._add_intro_text("Emerging risks and developments requiring attention.")

        if risk_radar:
            self._add_formatted_text(risk_radar)
        else:
            # Auto-generate risk indicators from high-importance articles
            high_risk = [a for a in articles if a.importance_score >= 7]
            if high_risk:
                self._add_sub_heading("HIGH-PRIORITY WATCH ITEMS")
                for article in sorted(high_risk, key=lambda x: x.importance_score, reverse=True)[:5]:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_after = Pt(4)

                    # Priority indicator
                    if article.importance_score >= 8:
                        tag = p.add_run("[CRITICAL] ")
                        tag.font.bold = True
                        tag.font.color.rgb = RED
                        tag.font.size = CAPTION_SIZE
                    else:
                        tag = p.add_run("[WATCH] ")
                        tag.font.bold = True
                        tag.font.color.rgb = ORANGE
                        tag.font.size = CAPTION_SIZE

                    title_run = p.add_run(f"{article.title}")
                    title_run.font.size = BODY_SIZE

                    source_run = p.add_run(f" — {article.source}")
                    source_run.font.size = SMALL_SIZE
                    source_run.font.color.rgb = GREY_TEXT

                    if article.ai_summary:
                        summary_p = self.doc.add_paragraph()
                        summary_p.paragraph_format.left_indent = Inches(0.4)
                        summary_p.paragraph_format.space_after = Pt(6)
                        run = summary_p.add_run(article.ai_summary)
                        run.font.size = SMALL_SIZE
            else:
                p = self.doc.add_paragraph()
                run = p.add_run("No high-priority risk items identified this week.")
                run.font.italic = True
                run.font.color.rgb = GREY_TEXT

    # =====================================================================
    # SECTION 9: THE WEEK AHEAD
    # =====================================================================

    def _add_week_ahead(self, week_ahead: str):
        """Add forward-looking section."""
        self._add_intro_text("Key events, data releases, and developments to watch in the coming week.")

        if week_ahead:
            self._add_formatted_text(week_ahead)
        else:
            self._add_sub_heading("WHAT TO WATCH")
            items = [
                "Monitor central bank communications for policy signals",
                "Track upcoming economic data releases (GDP, CPI, employment)",
                "Watch for geopolitical developments affecting trade and markets",
                "Review corporate earnings for sector-level economic signals",
            ]
            for item in items:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(3)
                self._add_inline_formatted_run(p, f"• {item}")

    # =====================================================================
    # SECTION 10: APPENDIX - COMPACT ARTICLE INDEX
    # =====================================================================

    def _add_article_appendix(self, all_articles: list[Article], articles_by_org: dict):
        """Add compact article reference table grouped by category."""
        self._add_intro_text(
            f"Complete index of {len(all_articles)} articles collected this week, "
            f"from {len(articles_by_org)} organizations."
        )

        # Group by category
        orgs_by_category = self._group_by_category(articles_by_org)

        for category in self._get_category_order():
            if category not in orgs_by_category:
                continue

            orgs = orgs_by_category[category]
            category_titles = {
                'Central Bank': 'Central Banks',
                'International': 'International & Multilateral Organizations',
                'Consulting': 'Consulting Firms & Professional Services',
                'Think Tank': 'Think Tanks & Research Institutes',
                'Investment Bank': 'Investment Banks Research',
                'News': 'News & Data Providers',
                'Data Provider': 'Data & Analytics Providers',
                'Rating Agency': 'Credit Rating Agencies',
                'India': 'India-Specific Sources',
            }

            title = category_titles.get(category, category)
            total = sum(len(arts) for arts in orgs.values())
            self._add_sub_heading(f"{title} ({total} articles)")

            # Create compact table for each category
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Header
            headers = ['Source', 'Title', 'Importance', 'AI Summary']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = WHITE
                        run.font.size = Pt(8)
                        run.font.name = FONT_FAMILY
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SHADE_TITLE_BG}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            # Articles
            for org_name in sorted(orgs.keys()):
                for article in orgs[org_name]:
                    row = table.add_row()
                    row.cells[0].text = article.source
                    row.cells[1].text = article.title[:80]
                    row.cells[2].text = f"{article.importance_score}/10"

                    # Importance color
                    if article.importance_score >= 8:
                        shade = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SHADE_CRITICAL}"/>')
                        row.cells[2]._tc.get_or_add_tcPr().append(shade)
                    elif article.importance_score >= 5:
                        shade = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SHADE_IMPORTANT}"/>')
                        row.cells[2]._tc.get_or_add_tcPr().append(shade)

                    summary = article.ai_summary or article.summary
                    row.cells[3].text = (summary[:120] + "...") if summary and len(summary) > 120 else (summary or "—")

                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                run.font.name = FONT_FAMILY

            # Column widths
            widths = [Inches(0.8), Inches(2.0), Inches(0.7), Inches(3.5)]
            for i, width in enumerate(widths):
                for row in table.rows:
                    row.cells[i].width = width

            self.doc.add_paragraph()  # Spacer between categories

    # =====================================================================
    # METHODOLOGY NOTE
    # =====================================================================

    def _add_methodology_note(self):
        """Add methodology and disclaimer."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        run = p.add_run("_" * 60)
        run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
        run.font.size = Pt(6)

        self._add_sub_heading("METHODOLOGY")
        method_text = (
            "This report is generated by the Economic Intelligence Agent using NVIDIA NIM's "
            "7-model architecture. Articles are collected from 39 leading organizations worldwide, "
            "deduplicated using semantic embeddings (NV-Embed-V1), ranked by economic relevance "
            "(Llama Nemotron Rerank), summarized (Mistral Small 3.1), deep-analyzed for major reports "
            "(DeepSeek V3.1 Terminus), and synthesized (Kimi K2 Instruct). PDF reports are processed "
            "for charts, tables, and key statistics."
        )
        p = self.doc.add_paragraph()
        run = p.add_run(method_text)
        run.font.size = CAPTION_SIZE
        run.font.color.rgb = GREY_TEXT

        self._add_sub_heading("DISCLAIMER")
        disclaimer = (
            "This report is generated by an AI system and is intended for informational purposes only. "
            "It does not constitute financial, investment, or policy advice. The analysis reflects "
            "AI-generated interpretations of publicly available information and may contain errors or "
            "omissions. Always verify critical data points from primary sources before making decisions."
        )
        p = self.doc.add_paragraph()
        run = p.add_run(disclaimer)
        run.font.size = CAPTION_SIZE
        run.font.color.rgb = GREY_TEXT
        run.font.italic = True

    # =====================================================================
    # TL;DR TOP 5
    # =====================================================================

    def _add_tldr_content(self, content: str):
        """Add TL;DR Top 5 with color-coded priority indicators."""
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue

            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Inches(0.2)

            if any(marker in stripped for marker in ['[CRITICAL]', '\U0001F534']):
                clean = stripped.replace('\U0001F534', '').strip()
                tag = p.add_run("[CRITICAL] ")
                tag.font.bold = True
                tag.font.color.rgb = RED
                tag.font.size = BODY_SIZE
                self._add_inline_formatted_run(p, clean)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SHADE_CRITICAL}"/>')
                p.paragraph_format.element.get_or_add_pPr().append(shading)

            elif any(marker in stripped for marker in ['[IMPORTANT]', '\U0001F7E0']):
                clean = stripped.replace('\U0001F7E0', '').strip()
                tag = p.add_run("[IMPORTANT] ")
                tag.font.bold = True
                tag.font.color.rgb = ORANGE
                tag.font.size = BODY_SIZE
                self._add_inline_formatted_run(p, clean)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SHADE_IMPORTANT}"/>')
                p.paragraph_format.element.get_or_add_pPr().append(shading)

            elif any(marker in stripped for marker in ['[NOTABLE]', '\U0001F7E2']):
                clean = stripped.replace('\U0001F7E2', '').strip()
                tag = p.add_run("[NOTABLE] ")
                tag.font.bold = True
                tag.font.color.rgb = GREEN
                tag.font.size = BODY_SIZE
                self._add_inline_formatted_run(p, clean)
            else:
                self._add_inline_formatted_run(p, stripped)

    # =====================================================================
    # SENTIMENT
    # =====================================================================

    def _add_sentiment_content(self, content: str):
        """Add sentiment analysis with color-coded indicators."""
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue

            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

            if 'BULLISH' in stripped.upper():
                clean = stripped.replace('\U0001F7E2', '[BULLISH]')
                run = p.add_run(clean)
                run.font.bold = True
                run.font.color.rgb = GREEN
            elif 'BEARISH' in stripped.upper():
                clean = stripped.replace('\U0001F534', '[BEARISH]')
                run = p.add_run(clean)
                run.font.bold = True
                run.font.color.rgb = RED
            elif 'NEUTRAL' in stripped.upper():
                clean = stripped.replace('\U0001F7E1', '[NEUTRAL]')
                run = p.add_run(clean)
                run.font.bold = True
                run.font.color.rgb = ORANGE
            elif stripped.startswith('**'):
                self._add_inline_formatted_run(p, stripped)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                p.paragraph_format.left_indent = Inches(0.3)
                self._add_inline_formatted_run(p, "  " + stripped[2:])
            else:
                self._add_inline_formatted_run(p, stripped)

    # =====================================================================
    # THEME CONTENT (fallback for Deep Dives)
    # =====================================================================

    def _add_theme_content(self, theme_summary: dict):
        """Add theme-based summary with clean formatting."""
        if not theme_summary:
            return

        for theme, data in theme_summary.items():
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f">> {theme}")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_BLUE

            stats = self.doc.add_paragraph()
            stats.paragraph_format.left_indent = Inches(0.3)
            stats.paragraph_format.space_after = Pt(2)
            run = stats.add_run(
                f"{data['count']} items from {len(data['sources'])} sources: "
                + ", ".join(data['sources'][:5])
            )
            run.font.size = SMALL_SIZE
            run.font.color.rgb = GREY_TEXT

            for headline in data.get('headlines', [])[:3]:
                hl = self.doc.add_paragraph()
                hl.paragraph_format.left_indent = Inches(0.5)
                hl.paragraph_format.space_after = Pt(1)
                run = hl.add_run(f"• {headline[:80]}")
                run.font.size = SMALL_SIZE

    # =====================================================================
    # MARKDOWN-TO-WORD PARSER
    # =====================================================================

    def _add_formatted_text(self, text: str):
        """Parse markdown-like text and render it properly in Word."""
        if not text:
            return

        paragraphs = text.split('\n')

        for line in paragraphs:
            stripped = line.strip()
            if not stripped:
                continue

            # Heading lines: **SOME HEADER** or ## Header
            if (stripped.startswith('**') and stripped.endswith('**') and
                    stripped.count('**') == 2):
                inner = stripped[2:-2].strip()
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(inner)
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_BLUE
                continue

            if stripped.startswith('## '):
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run(stripped[3:])
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_BLUE
                continue

            # Bullet points
            if stripped.startswith('- ') or stripped.startswith('* '):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(1)
                self._add_inline_formatted_run(p, "  " + stripped[2:])
                continue

            # Numbered items
            num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if num_match:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(3)
                num_run = p.add_run(f"{num_match.group(1)}. ")
                num_run.font.bold = True
                self._add_inline_formatted_run(p, num_match.group(2))
                continue

            # Regular paragraph
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            self._add_inline_formatted_run(p, stripped)

    def _add_inline_formatted_run(self, paragraph, text: str):
        """Parse inline **bold** and *italic* markers."""
        parts = re.split(r'(\*\*.*?\*\*)', text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                inner = part[2:-2]
                italic_parts = re.split(r'(\*.*?\*)', inner)
                for ip in italic_parts:
                    if ip.startswith('*') and ip.endswith('*'):
                        run = paragraph.add_run(ip[1:-1])
                        run.font.bold = True
                        run.font.italic = True
                    else:
                        run = paragraph.add_run(ip)
                        run.font.bold = True
            else:
                italic_parts = re.split(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', part)
                for i, ip in enumerate(italic_parts):
                    run = paragraph.add_run(ip)
                    if i % 2 == 1:
                        run.font.italic = True

    # =====================================================================
    # TABLE RENDERING
    # =====================================================================

    def _render_markdown_table(self, markdown_text: str):
        """Convert a markdown table to a Word table."""
        lines = [l.strip() for l in markdown_text.split('\n') if l.strip()]

        data_lines = []
        for line in lines:
            if line.startswith('|') and not re.match(r'^\|[\s\-:]+\|', line):
                cells = [c.strip() for c in line.strip('|').split('|')]
                data_lines.append(cells)
            elif not line.startswith('|'):
                if line.startswith('**'):
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    self._add_inline_formatted_run(p, line)
                    for r in p.runs:
                        r.font.size = SMALL_SIZE
                    return

        if len(data_lines) < 2:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(markdown_text[:300])
            run.font.size = SMALL_SIZE
            return

        num_cols = len(data_lines[0])
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'

        for i, cell_text in enumerate(data_lines[0]):
            cell = table.rows[0].cells[i]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = FONT_FAMILY

        for row_data in data_lines[1:]:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                if i < num_cols:
                    row.cells[i].text = cell_text
                    for paragraph in row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.name = FONT_FAMILY

        self.doc.add_paragraph()

    # =====================================================================
    # HELPERS
    # =====================================================================

    def _group_by_category(self, articles_by_org: dict) -> dict:
        """Group organizations by their category."""
        orgs_by_category = {}
        for org_name, articles in articles_by_org.items():
            if not articles:
                continue
            category = articles[0].category
            if category not in orgs_by_category:
                orgs_by_category[category] = {}
            orgs_by_category[category][org_name] = articles
        return orgs_by_category

    def _get_category_order(self) -> list[str]:
        return [
            'Central Bank', 'International', 'Consulting', 'Think Tank',
            'Investment Bank', 'News', 'Data Provider', 'Rating Agency', 'India',
        ]

    # =====================================================================
    # HYPERLINKS
    # =====================================================================

    def _add_hyperlink(self, paragraph, url: str, text: str = None):
        """Add a real clickable hyperlink to a paragraph."""
        if text is None:
            text = url

        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<w:r>'
            f'<w:rPr>'
            f'<w:rStyle w:val="Hyperlink"/>'
            f'<w:color w:val="0563C1"/>'
            f'<w:u w:val="single"/>'
            f'<w:sz w:val="{int(CAPTION_SIZE.pt * 2)}"/>'
            f'<w:rFonts w:ascii="{FONT_FAMILY}" w:hAnsi="{FONT_FAMILY}"/>'
            f'</w:rPr>'
            f'<w:t>{self._escape_xml(text)}</w:t>'
            f'</w:r>'
            f'</w:hyperlink>'
        )

        paragraph._p.append(hyperlink)

    @staticmethod
    def _escape_xml(text: str) -> str:
        """Escape special characters for XML."""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&apos;'))
